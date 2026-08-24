#!/usr/bin/env python3
"""Build Rhino MCP's local page-level regulation search database."""

from __future__ import annotations

import argparse
import concurrent.futures
import hashlib
import json
import re
import shutil
import sqlite3
import subprocess
import tempfile
import unicodedata
from collections import Counter
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from docx import Document
from PIL import Image
from pypdf import PdfReader

SCHEMA_VERSION = "1"
TEXT_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/html",
    "image/png",
    "image/jpeg",
    "image/tiff",
}
MOJIBAKE_CHARACTERS = frozenset(
    "ÃÂÐÑÞßàáâãäåæèéêëìíîïðñòóôõøùúûýþÿ"
)


class _HTMLText(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.parts.append(data)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--manifest", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--report", type=Path)
    parser.add_argument("--ocr", action="store_true", help="OCR image-only pages and images")
    parser.add_argument("--ocr-dpi", type=int, default=220)
    parser.add_argument("--ocr-timeout", type=int, default=180)
    parser.add_argument("--ocr-workers", type=int, default=4)
    return parser.parse_args()


def normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFC", value).replace("\x00", "")
    lines = []
    for line in value.replace("\r", "\n").split("\n"):
        line = re.sub(r"[\t \u00a0]+", " ", line).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def looks_like_html(path: Path) -> bool:
    with path.open("rb") as stream:
        return stream.read(32).lstrip().startswith(b"<")


def mojibake_ratio(text: str) -> float:
    letters = [character for character in text if character.isalpha()]
    if not letters:
        return 0.0
    suspicious = sum(character in MOJIBAKE_CHARACTERS for character in letters)
    return suspicious / len(letters)


def needs_ocr(text: str) -> bool:
    return len(re.findall(r"\w", text)) < 80 or mojibake_ratio(text) >= 0.06


def language(value: str) -> str:
    sample = value[:20_000]
    has_cyrillic = bool(re.search(r"[\u0400-\u04ff]", sample))
    has_azerbaijani = bool(re.search(r"[əƏğıİöÖşŞçÇüÜ]", sample))
    if has_cyrillic and has_azerbaijani:
        return "az+ru"
    if has_cyrillic:
        return "ru"
    if has_azerbaijani:
        return "az"
    return "unknown"


def category(folder: str, title: str, mime_type: str) -> str:
    folded = f"{folder} {title}".casefold()
    if "роза ветров" in folded or "roza vetrov" in folded:
        return "climate_reference"
    if folder == "Tehsil":
        return "education"
    if (
        folder.endswith("_files")
        or mime_type in {"text/css", "application/octet-stream"}
        or title.endswith(".download")
    ):
        return "support_file"
    if mime_type.startswith("image/") or mime_type == "image/vnd.dwg":
        return "drawing_reference"
    return "regulation"


def available_ocr_languages() -> str | None:
    if shutil.which("tesseract") is None:
        return None
    result = subprocess.run(
        ["tesseract", "--list-langs"], text=True, capture_output=True, check=False
    )
    installed = set(result.stdout.splitlines()[1:])
    selected = [name for name in ("aze", "rus", "eng") if name in installed]
    return "+".join(selected) if selected else None


def ocr_image(path: Path, languages: str, timeout: int, page_segmentation: int) -> str:
    result = subprocess.run(
        [
            "tesseract",
            str(path),
            "stdout",
            "-l",
            languages,
            "--psm",
            str(page_segmentation),
        ],
        text=True,
        capture_output=True,
        check=False,
        timeout=timeout,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or f"tesseract exited {result.returncode}")
    return normalize_text(result.stdout)


def extract_pdf(
    path: Path,
    *,
    use_ocr: bool,
    ocr_languages: str | None,
    ocr_dpi: int,
    ocr_timeout: int,
    ocr_workers: int,
) -> tuple[list[tuple[int, str, str]], int, list[str]]:
    reader = PdfReader(path)
    if reader.is_encrypted:
        reader.decrypt("")
    page_results: list[list[Any]] = []
    errors: list[str] = []
    poppler_pages: list[str] = []
    if shutil.which("pdftotext"):
        result = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(path), "-"],
            capture_output=True,
            check=False,
            timeout=max(ocr_timeout, 600),
        )
        if result.returncode == 0:
            poppler_pages = result.stdout.decode("utf-8", errors="replace").split("\f")
        else:
            errors.append(
                "Poppler text extraction: "
                + result.stderr.decode("utf-8", errors="replace").strip()
            )
    with tempfile.TemporaryDirectory(prefix="rhino-mcp-ocr-") as temp:
        temp_dir = Path(temp)
        for number, page in enumerate(reader.pages, 1):
            text = (
                normalize_text(poppler_pages[number - 1])
                if number <= len(poppler_pages)
                else ""
            )
            method = "pdf_poppler" if text else "pdf_text"
            if len(re.findall(r"\w", text)) < 80:
                try:
                    pypdf_text = normalize_text(page.extract_text() or "")
                    if len(pypdf_text) > len(text):
                        text = pypdf_text
                        method = "pdf_text"
                except Exception as exc:
                    if not text:
                        errors.append(f"page {number} text: {exc}")
            page_results.append([number, text, method])

        candidates = [
            number
            for number, text, _ in page_results
            if use_ocr and ocr_languages and needs_ocr(text)
        ]

        def run_ocr(number: int) -> tuple[int, str, str | None]:
            prefix = temp_dir / f"page-{number}"
            rendered = prefix.with_suffix(".png")
            command = [
                "pdftoppm",
                "-f",
                str(number),
                "-l",
                str(number),
                "-r",
                str(ocr_dpi),
                "-png",
                "-singlefile",
                str(path),
                str(prefix),
            ]
            try:
                render = subprocess.run(
                    command, text=True, capture_output=True, check=False, timeout=ocr_timeout
                )
                if render.returncode != 0:
                    raise RuntimeError(render.stderr.strip())
                return (
                    number,
                    ocr_image(rendered, ocr_languages, ocr_timeout, 3),
                    None,
                )
            except Exception as exc:
                return number, "", str(exc)
            finally:
                rendered.unlink(missing_ok=True)

        with concurrent.futures.ThreadPoolExecutor(
            max_workers=max(1, min(8, ocr_workers))
        ) as executor:
            for number, ocr_text, error in executor.map(run_ocr, candidates):
                if error:
                    errors.append(f"page {number} OCR: {error}")
                else:
                    original = page_results[number - 1][1]
                    better_encoding = (
                        mojibake_ratio(original) >= 0.06
                        and mojibake_ratio(ocr_text) < mojibake_ratio(original) / 2
                    )
                    if len(ocr_text) > len(original) or (len(ocr_text) >= 40 and better_encoding):
                        page_results[number - 1][1] = ocr_text
                        page_results[number - 1][2] = "pdf_ocr"

    pages = [tuple(result) for result in page_results if result[1]]
    return pages, len(reader.pages), errors


def extract_docx(path: Path) -> tuple[list[tuple[int, str, str]], int, list[str]]:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = normalize_text("\n".join(parts))
    return ([(1, text, "docx_text")] if text else []), 1, []


def extract_html(path: Path) -> tuple[list[tuple[int, str, str]], int, list[str]]:
    parser = _HTMLText()
    parser.feed(path.read_text(encoding="utf-8", errors="replace"))
    text = normalize_text("\n".join(parser.parts))
    return ([(1, text, "html_text")] if text else []), 1, []


def extract_image(
    path: Path, *, use_ocr: bool, ocr_languages: str | None, ocr_timeout: int
) -> tuple[list[tuple[int, str, str]], int, list[str]]:
    errors: list[str] = []
    pages: list[tuple[int, str, str]] = []
    with Image.open(path) as image:
        frames = int(getattr(image, "n_frames", 1))
        if not use_ocr or not ocr_languages:
            return pages, frames, errors
        with tempfile.TemporaryDirectory(prefix="rhino-mcp-image-ocr-") as temp:
            temp_dir = Path(temp)
            for number in range(1, frames + 1):
                try:
                    image.seek(number - 1)
                    frame = image.convert("RGB")
                    target = temp_dir / f"frame-{number}.png"
                    frame.save(target)
                    text = ocr_image(target, ocr_languages, ocr_timeout, 11)
                    if text:
                        pages.append((number, text, "image_ocr"))
                except Exception as exc:
                    errors.append(f"frame {number} OCR: {exc}")
    return pages, frames, errors


def extract(
    path: Path,
    mime_type: str,
    *,
    use_ocr: bool,
    ocr_languages: str | None,
    ocr_dpi: int,
    ocr_timeout: int,
    ocr_workers: int,
) -> tuple[list[tuple[int, str, str]], int, list[str]]:
    if mime_type == "application/pdf":
        if looks_like_html(path):
            return extract_html(path)
        return extract_pdf(
            path,
            use_ocr=use_ocr,
            ocr_languages=ocr_languages,
            ocr_dpi=ocr_dpi,
            ocr_timeout=ocr_timeout,
            ocr_workers=ocr_workers,
        )
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_docx(path)
    if mime_type == "text/html":
        return extract_html(path)
    if mime_type.startswith("image/") and mime_type != "image/vnd.dwg":
        return extract_image(
            path,
            use_ocr=use_ocr,
            ocr_languages=ocr_languages,
            ocr_timeout=ocr_timeout,
        )
    return [], 0, []


def create_schema(connection: sqlite3.Connection) -> None:
    connection.executescript(
        """
        PRAGMA journal_mode = OFF;
        PRAGMA synchronous = OFF;
        PRAGMA temp_store = MEMORY;
        CREATE TABLE metadata (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL
        );
        CREATE TABLE documents (
            id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            folder TEXT NOT NULL,
            category TEXT NOT NULL,
            language TEXT NOT NULL,
            mime_type TEXT NOT NULL,
            size_bytes INTEGER NOT NULL,
            modified_time TEXT,
            drive_url TEXT NOT NULL,
            local_path TEXT NOT NULL,
            total_pages INTEGER NOT NULL,
            indexed_pages INTEGER NOT NULL,
            character_count INTEGER NOT NULL,
            status TEXT NOT NULL,
            extraction_methods TEXT NOT NULL,
            errors TEXT NOT NULL
        );
        CREATE TABLE pages (
            id INTEGER PRIMARY KEY,
            document_id TEXT NOT NULL REFERENCES documents(id),
            page_number INTEGER NOT NULL,
            text TEXT NOT NULL,
            extraction_method TEXT NOT NULL,
            UNIQUE(document_id, page_number)
        );
        CREATE INDEX pages_document_page ON pages(document_id, page_number);
        CREATE VIRTUAL TABLE pages_fts USING fts5(
            title,
            text,
            content='',
            tokenize='unicode61 remove_diacritics 2'
        );
        """
    )


def build(args: argparse.Namespace) -> dict[str, Any]:
    manifest_path = args.manifest.resolve()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    source_root = manifest_path.parent
    output = args.output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_suffix(output.suffix + ".building")
    temporary.unlink(missing_ok=True)
    ocr_languages = available_ocr_languages() if args.ocr else None
    if args.ocr and not ocr_languages:
        raise RuntimeError("--ocr requires tesseract with at least one installed language.")
    if args.ocr and shutil.which("pdftoppm") is None:
        raise RuntimeError("--ocr requires pdftoppm for image-only PDF pages.")

    connection = sqlite3.connect(temporary)
    create_schema(connection)
    generated = datetime.now(timezone.utc).isoformat()
    manifest_hash = hashlib.sha256(manifest_path.read_bytes()).hexdigest()
    metadata = {
        "schema_version": SCHEMA_VERSION,
        "generated_at": generated,
        "snapshot_created": manifest.get("generated_at", ""),
        "source_folder_url": manifest.get("source_folder_url", ""),
        "source_file_count": str(manifest.get("file_count", len(manifest["files"]))),
        "source_total_bytes": str(manifest.get("total_bytes", 0)),
        "manifest_sha256": manifest_hash,
        "ocr_languages": ocr_languages or "disabled",
        "notice": "Reference snapshot only; verify applicability and current legal status.",
    }
    connection.executemany("INSERT INTO metadata(key, value) VALUES (?, ?)", metadata.items())

    methods: Counter[str] = Counter()
    failures: list[dict[str, Any]] = []
    indexed_documents = 0
    total_indexed_pages = 0
    total_characters = 0
    for position, item in enumerate(manifest["files"], 1):
        path = source_root / item["local_path"]
        errors: list[str] = []
        pages: list[tuple[int, str, str]] = []
        total_pages = 0
        if not path.is_file():
            errors.append("source file is missing")
        elif path.stat().st_size != int(item["size"]):
            errors.append("source file size does not match the Drive manifest")
        elif item["mime_type"] in TEXT_MIME_TYPES:
            try:
                pages, total_pages, extraction_errors = extract(
                    path,
                    item["mime_type"],
                    use_ocr=args.ocr,
                    ocr_languages=ocr_languages,
                    ocr_dpi=args.ocr_dpi,
                    ocr_timeout=args.ocr_timeout,
                    ocr_workers=args.ocr_workers,
                )
                errors.extend(extraction_errors)
            except Exception as exc:
                errors.append(str(exc))
        text_sample = "\n".join(text for _, text, _ in pages)
        char_count = sum(len(text) for _, text, _ in pages)
        status = "indexed" if pages else ("error" if errors else "not_textual")
        extraction_methods = sorted({method for _, _, method in pages})
        connection.execute(
            """
            INSERT INTO documents(
                id, title, folder, category, language, mime_type, size_bytes,
                modified_time, drive_url, local_path, total_pages, indexed_pages,
                character_count, status, extraction_methods, errors
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                item["id"],
                item["title"],
                item.get("relative_folder", ""),
                category(item.get("relative_folder", ""), item["title"], item["mime_type"]),
                language(f"{item['title']}\n{text_sample}"),
                item["mime_type"],
                int(item["size"]),
                item.get("modified_time"),
                item["drive_url"],
                item["local_path"],
                total_pages,
                len(pages),
                char_count,
                status,
                json.dumps(extraction_methods, ensure_ascii=False),
                json.dumps(errors, ensure_ascii=False),
            ),
        )
        for page_number, text, method in pages:
            cursor = connection.execute(
                "INSERT INTO pages(document_id, page_number, text, extraction_method) "
                "VALUES (?, ?, ?, ?)",
                (item["id"], page_number, text, method),
            )
            connection.execute(
                "INSERT INTO pages_fts(rowid, title, text) VALUES (?, ?, ?)",
                (cursor.lastrowid, item["title"], text),
            )
            methods[method] += 1
        indexed_documents += int(bool(pages))
        total_indexed_pages += len(pages)
        total_characters += char_count
        if errors:
            failures.append({"id": item["id"], "title": item["title"], "errors": errors})
        if position % 10 == 0 or position == len(manifest["files"]):
            print(
                f"[{position:>3}/{len(manifest['files'])}] "
                f"{indexed_documents} documents, {total_indexed_pages} pages",
                flush=True,
            )
            connection.commit()

    connection.execute("ANALYZE")
    connection.commit()
    connection.execute("VACUUM")
    connection.close()
    temporary.replace(output)

    report = {
        "generated_at": generated,
        "database": str(output),
        "database_bytes": output.stat().st_size,
        "source_files": len(manifest["files"]),
        "indexed_documents": indexed_documents,
        "indexed_pages": total_indexed_pages,
        "characters": total_characters,
        "extraction_methods": dict(sorted(methods.items())),
        "ocr_languages": ocr_languages,
        "documents_with_errors": len(failures),
        "errors": failures,
    }
    if args.report:
        args.report.resolve().parent.mkdir(parents=True, exist_ok=True)
        args.report.resolve().write_text(
            json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
        )
    return report


def main() -> int:
    args = parse_args()
    report = build(args)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
